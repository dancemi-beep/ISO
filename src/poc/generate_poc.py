import docx

def replace_text_in_doc(doc, data_dict):
    """將文件中的變數標籤替換為真實資料"""
    # 替換段落中的文字
    for para in doc.paragraphs:
        for key, value in data_dict.items():
            if key in para.text:
                # 簡單替換 (若變數橫跨多個 runs 則直接取代 Paragraph.text)
                inline = para.runs
                for i in range(len(inline)):
                    if key in inline[i].text:
                        inline[i].text = inline[i].text.replace(key, value)
                
                if key in para.text and not any(key in r.text for r in inline):
                    para.text = para.text.replace(key, value)

    # 替換表格中的文字
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    for key, value in data_dict.items():
                        if key in para.text:
                            inline = para.runs
                            for i in range(len(inline)):
                                if key in inline[i].text:
                                    inline[i].text = inline[i].text.replace(key, value)
                            
                            if key in para.text and not any(key in r.text for r in inline):
                                para.text = para.text.replace(key, value)

if __name__ == "__main__":
    template_path = "src/poc/IS-001_template_marked.docx"
    output_path = "src/poc/IS-001_output_demo.docx"
    
    # 模擬從問卷回答中取得的資料
    fake_data = {
        "{{company_name}}": "雲端守護科技有限公司",
        "{{publish_date}}": "2026/03/05"
    }
    
    print(f"開始讀取範本檔案: {template_path}")
    doc = docx.Document(template_path)
    
    print("正在將變數替換為真實資料...")
    replace_text_in_doc(doc, fake_data)
    
    doc.save(output_path)
    print(f"文件生成成功！已儲存至: {output_path}")
