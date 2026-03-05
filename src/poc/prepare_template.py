import docx
import re

def replace_text_in_doc(doc, old_text, new_text):
    for para in doc.paragraphs:
        if old_text in para.text:
            # Simple replace: might lose formatting if it spans across runs,
            # but for POC we replace text in the paragraph directly
            # A safer way is to clear runs and add new text if we just replace it
            inline = para.runs
            for i in range(len(inline)):
                if old_text in inline[i].text:
                    inline[i].text = inline[i].text.replace(old_text, new_text)
            
            # If not in a single run (often the case in Word docs), we fall back:
            if old_text in para.text and not any(old_text in r.text for r in inline):
                para.text = para.text.replace(old_text, new_text)

    # Also replace in tables
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    if old_text in para.text:
                        inline = para.runs
                        for i in range(len(inline)):
                            if old_text in inline[i].text:
                                inline[i].text = inline[i].text.replace(old_text, new_text)
                        if old_text in para.text and not any(old_text in r.text for r in inline):
                            para.text = para.text.replace(old_text, new_text)

if __name__ == "__main__":
    file_path = "src/poc/IS-001_template.docx"
    doc = docx.Document(file_path)
    
    # Simulate user preparing template by adding {{}} tags
    replace_text_in_doc(doc, "XXXX股份有限公司", "{{company_name}}")
    replace_text_in_doc(doc, "年/月/日", "{{publish_date}}")
    
    output_path = "src/poc/IS-001_template_marked.docx"
    doc.save(output_path)
    print(f"範本準備完成，已儲存至: {output_path}")
