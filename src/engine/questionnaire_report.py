import json
import os
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.style import WD_STYLE_TYPE

class QuestionnaireReportGenerator:
    def __init__(self, config_path):
        import json
        with open(config_path, 'r', encoding='utf-8') as f:
            self.config = json.load(f)

    def generate(self, user_data, output_path):
        doc = Document()
        
        # Add basic styling
        styles = doc.styles
        style = styles.add_style('QuestionStyle', WD_STYLE_TYPE.PARAGRAPH)
        style.font.name = 'Microsoft JhengHei'
        style.font.bold = True
        style.font.size = Pt(12)
        
        answer_style = styles.add_style('AnswerStyle', WD_STYLE_TYPE.PARAGRAPH)
        answer_style.font.name = 'Microsoft JhengHei'
        answer_style.font.size = Pt(11)

        # Document Title
        title = doc.add_heading('ISO 27001 ISMS 導入問卷填答總表', level=0)
        title.alignment = 1 # Center

        description = doc.add_paragraph("本報告記錄了貴公司在系統化自動生成 ISMS 文件前，所填寫的所有問卷設定與管理決策參數。")
        description.paragraph_format.space_after = Pt(20)

        # We assume the config has stages 1-3
        for stage_key in ["stage1", "stage2", "stage3"]:
            stage_data = self.config.get(stage_key, {})
            if not stage_data:
                continue

            stage_title = stage_data.get("title", f"Stage {stage_key[-1]}")
            heading = doc.add_heading(stage_title, level=1)
            heading.paragraph_format.space_before = Pt(20)
            heading.paragraph_format.space_after = Pt(10)

            # Iterating sections in the given stage
            sections = stage_data.get("sections", [])
            for section in sections:
                sec_title = section.get('title', '')
                if sec_title:
                    h2 = doc.add_heading(sec_title, level=2)
                    h2.paragraph_format.space_before = Pt(10)
                
                fields = section.get('fields', [])
                for field in fields:
                    var_name = field['var']
                    label = field['label']
                    value = user_data.get(var_name, "（未填寫）")
                    
                    if not value:
                         value = "（未填寫）"

                    # Print Question
                    p_q = doc.add_paragraph(style='QuestionStyle')
                    p_q.add_run(f"■ {label} ({var_name})")
                    
                    # Print Answer
                    p_a = doc.add_paragraph(style='AnswerStyle')
                    p_a.add_run(str(value))
                    p_a.paragraph_format.left_indent = Inches(0.2)
                    p_a.paragraph_format.space_after = Pt(10)
        
        doc.save(output_path)
        print(f"Generated Questionnaire Report at {output_path}")

if __name__ == '__main__':
    # Local Test
    gen = QuestionnaireReportGenerator("questionnaire_config.json")
    gen.generate({"company_name": "Test Co."}, "output/問卷填答總表.docx")
