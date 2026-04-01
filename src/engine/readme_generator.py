import os
import json
from datetime import datetime
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT


class ReadmeGenerator:
    """Generates a '導入指引手冊.docx' based on user data and document structure."""

    def __init__(self, structure_file):
        # Support encrypted structure file
        enc_structure = structure_file + ".enc"
        if os.path.exists(enc_structure):
            from security import SecurityEngine
            sec = SecurityEngine()
            data = sec.decrypt_to_memory(enc_structure)
            self.structure = json.load(data)
        else:
            with open(structure_file, 'r', encoding='utf-8') as f:
                self.structure = json.load(f)

    def generate(self, user_data, output_dir, generated_files=None):
        """Create a landing readme .docx file.
        
        Args:
            user_data: dict of user questionnaire answers
            output_dir: path to save the readme
            generated_files: optional list of successfully generated filenames
        """
        doc = Document()

        company = user_data.get('company_name', '[公司名稱]')
        date_str = user_data.get('publish_date', datetime.now().strftime('%Y/%m/%d'))

        # --- Title ---
        title = doc.add_heading(f'{company}', level=0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER

        subtitle = doc.add_heading('ISO 27001 ISMS 導入文件指引手冊', level=1)
        subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER

        doc.add_paragraph(f'產出日期：{date_str}').alignment = WD_ALIGN_PARAGRAPH.CENTER
        doc.add_paragraph('')

        # --- Section 1: Document List ---
        doc.add_heading('一、文件清單一覽', level=2)
        doc.add_paragraph(
            '下表列出本次系統自動產出之全部 ISMS 文件，依編碼與階層分類。'
        )

        # Build table
        table = doc.add_table(rows=1, cols=4)
        table.style = 'Table Grid'
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        headers = ['文件編碼', '文件名稱', '階層', '類型']
        for i, h in enumerate(headers):
            cell = table.rows[0].cells[i]
            cell.text = h
            for para in cell.paragraphs:
                for run in para.runs:
                    run.bold = True
                    run.font.size = Pt(10)

        file_count = 0
        for family_key, files in self.structure.items():
            if family_key == "Other":
                continue
            for file_info in files:
                fname = file_info['filename']
                # If we have a generated list, only show files that were actually generated
                if generated_files and fname not in generated_files:
                    continue
                tier = file_info.get('tier', '-')
                ftype = 'Word' if file_info.get('type') == 'word' else 'Excel'
                row = table.add_row().cells
                row[0].text = family_key
                row[1].text = fname
                row[2].text = tier
                row[3].text = ftype
                file_count += 1

        doc.add_paragraph(f'\n共計 {file_count} 份文件。')

        # --- Section 2: Signing Process ---
        doc.add_heading('二、建議簽核流程', level=2)
        steps = [
            '1. 由資訊安全管理代表 (ISMS Representative) 進行文件初審。',
            '2. 各權責單位主管審閱所轄程序書與作業指導書。',
            '3. 由最高管理階層（如總經理）核准一階文件（資訊安全政策、適用性聲明書）。',
            '4. 核准後統一發行，建立「管制文件一覽表 (IS-002-04)」追蹤版本。',
            '5. 全體同仁進行資訊安全政策宣導與簽署確認。',
        ]
        for step in steps:
            doc.add_paragraph(step)

        # --- Section 3: Ops Reminders ---
        doc.add_heading('三、後續維運提醒', level=2)
        reminders = [
            '• 每年至少一次內部稽核 (參照 IS-002-07 內部稽核計畫)。',
            '• 每年至少一次管理審查會議 (參照 IS-002-09 管理審查會議紀錄)。',
            '• 資產及風險評鑑表 (IS-004-01) 需於組織重大變更時重新評估。',
            '• 教育訓練至少每年舉辦一次 (參照 IS-003-02 教育訓練計畫表)。',
            '• 營運持續計畫演練至少每年一次 (參照 IS-013-02)。',
            '• 所有表單紀錄保存期限建議至少三年。',
        ]
        for r in reminders:
            doc.add_paragraph(r)

        # --- Save ---
        output_path = os.path.join(output_dir, '導入指引手冊.docx')
        doc.save(output_path)
        print(f"Generated readme: {output_path}")
        return output_path


if __name__ == "__main__":
    gen = ReadmeGenerator(structure_file="document_structure.json")
    test_data = {
        "company_name": "新世紀虛擬科技股份有限公司",
        "publish_date": "2026/05/20",
    }
    gen.generate(test_data, output_dir="output")
