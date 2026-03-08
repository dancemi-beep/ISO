import os
import re
import docx


def replace_text_in_doc_runs_safe(doc, old_text, new_text):
    """
    Replace text in paragraphs, tables, headers, and footers.
    For preparation (Jinja tagging), para.text wholesale replacement 
    is acceptable since docxtpl needs {{ var }} in a single run.
    """
    for para in doc.paragraphs:
        if old_text in para.text:
            para.text = para.text.replace(old_text, new_text)

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    if old_text in para.text:
                        para.text = para.text.replace(old_text, new_text)

    for section in doc.sections:
        for para in section.header.paragraphs:
            if old_text in para.text:
                para.text = para.text.replace(old_text, new_text)
        for para in section.footer.paragraphs:
            if old_text in para.text:
                para.text = para.text.replace(old_text, new_text)


def batch_mark_templates(directory):
    # Replacement rules: ordered from most specific to least specific
    # to avoid partial matches (e.g. "OOOO股份有限公司" before "OOOO")
    replacements = [
        # Full company name patterns
        ("XXXX股份有限公司", "{{company_name}}"),
        ("OOOO股份有限公司", "{{company_name}}"),
        ("XXX股份有限公司", "{{company_name}}"),
        ("○○○○○股份有限公司", "{{company_name}}"),
        # Record number patterns (keep template format)
        ("IS-R-26-XXXXXX-XX", "IS-R-{{publish_date_short}}-XXXXXX-XX"),
        # Date patterns
        ("年/月/日", "{{publish_date}}"),
        ("YYYY/MM/DD", "{{publish_date}}"),
        ("OOO年", "{{publish_date}}"),
        # Jurisdiction
        ("XXX地方法院", "{{jurisdiction}}"),
        # URL patterns
        ("https://XXX.XXX.tw", "https://{{company_domain}}"),
        # Department placeholders in lists (e.g. "XXX部門")
        ("XXX部門", "{{dept_name}}部門"),
        # Specific P2/P3 Policy Replacements
        ("本公司所有同仁因工作目的，將個人自攜資訊設備（如筆記型電腦、智能手機、平板電腦等）連接公司網路作業工作環境時，必須填寫「員工攜帶自有設備申請書」，先行取得公司授權及同意。", "{{byod_policy_clause}}"),
        ("自攜資訊設備作業系統必須開啟自動更新，保持更新至最新版本；並安裝適切防毒軟體與及時更新病毒碼。", "{{byod_security_clause}}"),
        ("為求落實安全管理，於防火牆安全規則建置後，應每年審查規則之適用性，並予以檢討修正，避免因時間或業務變更而不符現狀。", "{{firewall_review_clause}}"),
        ("網頁過濾：組織應定期檢視相關網路設定，如需變更相關規則，應填寫「資訊服務申請單」，並針對員工對外部資源存取行為進行過濾及審查，避免遭受惡意攻擊，可考量下列方式以強化對外網站存取行為之安全性：", "{{web_filtering_clause}}"),
        ("應建立網路傳送資訊之安全控管機制（如防火牆、安全憑證、VPN或加解密系統或服務等），以確保資訊傳輸交換安全。", "{{remote_access_clause}}"),
        ("核心系統與設備之網路安全相關記錄檔（Log）案視資源使用狀況，以保存六個月為原則", "核心系統與設備之網路安全相關記錄檔（Log）案視資源使用狀況，以保存{{log_retention_months}}個月為原則"),
        ("發現疑似資訊安全異常事件或事故時，本公司同仁與委外人員皆負有即時通報之責任。", "{{incident_report_clause}}"),
        # Target passwords
        ("密碼長度最少12位長度", "密碼長度最少{{password_min_length}}位長度"),
        ("至少包含英文大寫、小寫字母、數字及特殊字元。", "{{password_complexity}}"),
        ("密碼最長使用期限為90天", "密碼最長使用期限為{{password_expiry_days}}天"),
        ("每90天更換通行密碼一次", "每{{password_expiry_days}}天更換通行密碼一次"),
        # IS-004-03 Risk Assessment overrides
        ("可接受風險值XX", "可接受風險值{{acceptable_risk_level}}"),
        ("建議值為XX", "建議值為{{acceptable_risk_level}}"),
        ("風險值XX", "風險值{{acceptable_risk_level}}"),
        # IS-001-02 Company Introduction
        ("Fabless IC設計公司，提供顯示器驅動DDIC (display driver IC)、電源控制晶片(Power IC)、光學感測器 (Optical sensors)、微機電感測器 (MEMS sensors)、電容觸控晶片(Touch controller IC)設計", "{{industry_description}}"),
        # Short name / remaining placeholders (must come AFTER longer patterns)
        ("OOOO", "{{company_short_name}}"),
        ("OOO", "{{company_short_name}}"),
        ("公司系統代碼", "{{company_short_name}}"),
    ]

    output_dir = os.path.join(directory, "marked")
    if not os.path.exists(output_dir):
        os.makedirs(output_dir)

    for f in sorted(os.listdir(directory)):
        if not f.endswith('.docx') or f.startswith('~'):
            continue

        filepath = os.path.join(directory, f)
        outpath = os.path.join(output_dir, f)

        try:
            doc = docx.Document(filepath)
            for old_str, jinja_var in replacements:
                replace_text_in_doc_runs_safe(doc, old_str, jinja_var)

            doc.save(outpath)
            print(f"Processed: {f}")
        except Exception as e:
            print(f"Error marking {f}: {e}")

    # Also copy xlsx files that don't need Jinja tagging
    import shutil
    for f in sorted(os.listdir(directory)):
        if f.endswith('.xlsx') and not f.startswith('~'):
            src = os.path.join(directory, f)
            dst = os.path.join(output_dir, f)
            shutil.copy2(src, dst)
            print(f"Copied xlsx: {f}")


if __name__ == "__main__":
    target_dir = "files"
    print(f"Starting batch marking in {target_dir}...")
    batch_mark_templates(target_dir)
    print("Done. Check files/marked/")
