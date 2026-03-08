import os
import json
import shutil
from docxtpl import DocxTemplate
import openpyxl


class DocumentGenerator:
    def __init__(self, template_dir, output_dir, structure_file):
        self.template_dir = template_dir
        self.output_dir = output_dir
        self.structure_file = structure_file

        with open(self.structure_file, 'r', encoding='utf-8') as f:
            self.structure = json.load(f)

        if not os.path.exists(self.output_dir):
            os.makedirs(self.output_dir)

    def generate_document(self, template_path, output_path, variables, user_data):
        """Generate a single .docx document by replacing variables using docxtpl."""
        try:
            doc = DocxTemplate(template_path)
            context = dict(user_data)
            
            doc_prefix = str(user_data.get("doc_prefix", "IS")).strip() or "IS"
            
            # Still provide [Missing: var] for any missing explicit variables
            for var in variables:
                if var not in context or not str(context[var]).strip():
                    context[var] = f"[Missing: {var}]"
            
            # 1. Render Jinja tags
            doc.render(context)
            
            # 2. Global Content Replacement (IS- -> PREFIX-)
            if doc_prefix != "IS":
                old_p = "IS-"
                new_p = f"{doc_prefix}-"
                
                # Helper to replace in paragraphs
                def replace_in_paras(paras):
                    for para in paras:
                        if old_p in para.text:
                            # Reconstruct text while trying to preserve some formatting if possible
                            # simple replacement on para.text might break runs, but docxtpl render already happened
                            # to be safe with runs:
                            for run in para.runs:
                                if old_p in run.text:
                                    run.text = run.text.replace(old_p, new_p)

                replace_in_paras(doc.paragraphs)
                
                for table in doc.tables:
                    for row in table.rows:
                        for cell in row.cells:
                            replace_in_paras(cell.paragraphs)
                
                for section in doc.sections:
                    for header in [section.header, section.first_page_header, section.even_page_header]:
                        if header: replace_in_paras(header.paragraphs)
                    for footer in [section.footer, section.first_page_footer, section.even_page_footer]:
                        if footer: replace_in_paras(footer.paragraphs)

            # 3. Apply Tier 1/2 Bold 24pt formatting to the company_name if present
            if "Tier 1" in template_path or "Tier 2" in template_path or "政策" in template_path or "程序" in template_path:
                from docx.shared import Pt
                for para in doc.paragraphs:
                    if user_data.get("company_name", "") in para.text:
                        for run in para.runs:
                            if user_data.get("company_name", "") in run.text:
                                run.font.bold = True
                                run.font.size = Pt(24)
                        break 

            doc.save(output_path)
            return True
        except Exception as e:
            print(f"Error processing {template_path}: {e}")
            return False

    def generate_excel(self, template_path, output_path, variables, user_data):
        """Generate a single .xlsx file by copying template and replacing cell values."""
        try:
            shutil.copy2(template_path, output_path)
            wb = openpyxl.load_workbook(output_path)
            
            doc_prefix = str(user_data.get("doc_prefix", "IS")).strip() or "IS"

            # Build replacement map from variables
            replacements = {}
            for var in variables:
                val = user_data.get(var)
                if val is not None:
                    replacements[var] = val

            # Scan all sheets for Jinja-like {{ var }} patterns and global IS- prefix
            for sheet in wb.sheetnames:
                ws = wb[sheet]
                for row in ws.iter_rows():
                    for cell in row:
                        if cell.value and isinstance(cell.value, str):
                            new_val = cell.value
                            # Replacement A: Variable Tags
                            for var, val in replacements.items():
                                new_val = new_val.replace("{{" + var + "}}", str(val))
                                new_val = new_val.replace("{{ " + var + " }}", str(val))
                            
                            # Replacement B: Global Prefix (IS- -> PREFIX-)
                            if doc_prefix != "IS":
                                new_val = new_val.replace("IS-", f"{doc_prefix}-")
                                
                            if new_val != cell.value:
                                cell.value = new_val

            # Special logic for IS-004-01 Asset & Risk Assessment
            if "IS-004-01" in template_path:
                ws = wb.active
                core_systems = user_data.get("core_systems", "")
                if core_systems:
                    # Parse systems (split by common delimiters)
                    import re
                    systems = [s.strip() for s in re.split(r'[、,\n]', core_systems) if s.strip()]
                    
                    # Start writing from row 5, col 2 (B5) for Asset Name
                    start_row = 5
                    for i, system in enumerate(systems):
                        row = start_row + i
                        ws.cell(row=row, column=2).value = system
                        ws.cell(row=row, column=3).value = f"自動帶入系統: {system}"
                        ws.cell(row=row, column=5).value = "軟體類"
                        ws.cell(row=row, column=6).value = "應用系統"
                        
                        # Set default CIA to 2 so formula logic works
                        ws.cell(row=row, column=9).value = 2 # C
                        ws.cell(row=row, column=10).value = 2 # I
                        ws.cell(row=row, column=11).value = 2 # A
                        
                        # Copy formulas from row 5 logic if we go past row 5
                        if row > 5:
                            ws.cell(row=row, column=12).value = f"=SUM(I{row}:K{row})" # P
                            ws.cell(row=row, column=25).value = f'=IF(L{row}*U{row}*W{row}*X{row}=0,"",L{row}*U{row}*W{row}*X{row})'
                            ws.cell(row=row, column=31).value = f'=IF(AA{row}*AB{row}*AC{row}*AD{row}=0,"",AA{row}*AB{row}*AC{row}*AD{row})'

            wb.save(output_path)
            return True
        except Exception as e:
            print(f"Error processing Excel {template_path}: {e}")
            return False

    def generate_all(self, user_data):
        """Iterate through the structure, auto-inject doc_number, and generate all documents.
        
        Returns:
            dict: {"success": [...], "failed": [...], "skipped": [...]}
        """
        result = {"success": [], "failed": [], "skipped": []}

        # --- Derive Policy Clauses from P2/P3 Questionnaire Choices ---
        # 1. BYOD Policy
        byod_choice = user_data.get("byod_policy", "")
        if byod_choice == "嚴格禁止自有設備處理公務":
            user_data["byod_policy_clause"] = "本公司嚴格禁止同仁將個人自攜資訊設備（如筆記型電腦、智能手機、平板電腦等）連接公司網路或處理公務資料。"
            user_data["byod_security_clause"] = "（因全面禁止 BYOD，故不適用自攜設備安全規範）"
        elif byod_choice == "自由開放":
            user_data["byod_policy_clause"] = "本公司開放同仁因工作目的，將個人自攜資訊設備連接公司網路或作業環境，惟仍須遵守基本安全規範。"
            user_data["byod_security_clause"] = "自攜資訊設備作業系統建議保持更新至最新版本；並安裝防毒軟體與及時更新病毒碼。"
        else: # Default: 有條件開放
            user_data["byod_policy_clause"] = "本公司所有同仁因工作目的，將個人自攜資訊設備（如筆記型電腦、智能手機、平板電腦等）連接公司網路作業工作環境時，必須填寫「員工攜帶自有設備申請書」，先行取得公司授權及同意。"
            user_data["byod_security_clause"] = "自攜資訊設備作業系統必須開啟自動更新，保持更新至最新版本；並安裝適切防毒軟體與及時更新病毒碼。"

        # 2. Remote Access Policy
        remote_choice = user_data.get("remote_access_policy", "")
        if remote_choice == "必須使用 VPN + MFA (多因素驗證)":
            user_data["remote_access_clause"] = "應建立網路傳送資訊之安全控管機制，且所有遠端連線均須透過虛擬私人網路 (VPN) 並啟用多重因素驗證 (MFA)，以確保資訊傳輸交換安全。"
        elif remote_choice == "無限制開放存取":
            user_data["remote_access_clause"] = "公司對外服務以便利性為優先，外部存取直接透過網際網路連線，惟重要系統仍應具備基本登入驗證。"
        else:
            user_data["remote_access_clause"] = "應建立網路傳送資訊之安全控管機制（如防火牆、安全憑證、VPN或加解密系統或服務等），以確保資訊傳輸交換安全。"

        # 3. Web Filtering Policy
        user_data["web_filtering_clause"] = "網頁過濾：組織應定期檢視相關網路設定，如需變更相關規則，應填寫「資訊服務申請單」，並針對員工對外部資源存取行為進行過濾及審查，避免遭受惡意攻擊，可考量下列方式以強化對外網站存取行為之安全性："
        user_data["firewall_review_clause"] = "為求落實安全管理，於防火牆安全規則建置後，應每年審查規則之適用性，並予以檢討修正，避免因時間或業務變更而不符現狀。"
        
        # 4. Incident Reporting
        incident_time = user_data.get("incident_report_time", "發現後 1 小時內")
        user_data["incident_report_clause"] = f"發現疑似資訊安全異常事件或事故時，本公司同仁與委外人員皆負有於{incident_time}即時通報之責任。"

        # 5. Log Retention (mapping select option to number)
        user_data["log_retention_months"] = "6"

        doc_prefix = str(user_data.get("doc_prefix", "IS")).strip()
        if not doc_prefix:
            doc_prefix = "IS"
            
        def apply_prefix(text):
            if not text: return text
            # Only replace the leading 'IS'
            if text.startswith("IS"):
                return text.replace("IS", doc_prefix, 1)
            return text

        for original_family_key, files in self.structure.items():
            family_key = apply_prefix(original_family_key)
            
            for file_info in files:
                original_filename = file_info['filename']
                filename = apply_prefix(original_filename)
                
                variables = file_info.get('variables', [])
                file_type = file_info.get('type', 'word')

                template_path = os.path.join(self.template_dir, original_filename)
                output_path = os.path.join(self.output_dir, filename)

                if not os.path.exists(template_path):
                    print(f"Warning: Template not found for {original_filename}")
                    result["skipped"].append({"name": filename, "reason": "範本檔案不存在"})
                    continue

                # --- Dynamic Exclusions based on Policies ---
                if "IS-008-04" in original_filename and user_data.get("byod_policy") == "嚴格禁止自有設備處理公務":
                    print(f"Skipping {filename} due to BYOD policy")
                    result["skipped"].append({"name": filename, "reason": "BYOD政策設定為嚴格禁止，無需此表單"})
                    continue
                
                # Auto-inject doc_number from family key (e.g. "ABC-001")
                enriched_data = dict(user_data)
                if 'doc_number' in variables and original_family_key != "Other":
                    enriched_data.setdefault('doc_number', family_key)

                if file_type == 'word' and original_filename.endswith('.docx'):
                    ok = self.generate_document(template_path, output_path, variables, enriched_data)
                elif file_type == 'excel' and original_filename.endswith('.xlsx'):
                    ok = self.generate_excel(template_path, output_path, variables, enriched_data)
                else:
                    print(f"Skipping unsupported type: {filename}")
                    result["skipped"].append({"name": filename, "reason": "系統不支援此檔案類型"})
                    continue

                if ok:
                    print(f"Generated: {filename}")
                    result["success"].append(filename)
                else:
                    result["failed"].append({"name": filename, "reason": "處理過程發生預期外的錯誤"})

        return result


if __name__ == "__main__":
    test_structure = "document_structure.json"
    generator = DocumentGenerator(
        template_dir="files/marked",
        output_dir="output",
        structure_file=test_structure
    )

    test_user_data = {
        "company_name": "新世紀虛擬科技股份有限公司",
        "company_short_name": "新世紀",
        "publish_date": "2026/05/20",
        "version": "1.0.0",
        "department": "資訊安全處"
    }

    print("Testing Batch Engine on entire marked folder...")
    result = generator.generate_all(test_user_data)
    print(f"\n--- Summary ---")
    print(f"Success: {len(result['success'])}")
    print(f"Failed:  {len(result['failed'])}")
    print(f"Skipped: {len(result['skipped'])}")
