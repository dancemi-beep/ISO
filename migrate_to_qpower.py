import os
import re
import json
import docx
import openpyxl

def replace_in_docx(filepath):
    doc = docx.Document(filepath)
    old_p = "IS-"
    new_p = "qpower-"
    
    def replace_in_paras(paras):
        for para in paras:
            if old_p in para.text:
                replaced = False
                for run in para.runs:
                    if old_p in run.text:
                        run.text = run.text.replace(old_p, new_p)
                        replaced = True
                if not replaced:
                    for i in range(len(para.runs) - 1):
                        if para.runs[i].text.endswith("IS") and para.runs[i+1].text.startswith("-"):
                            para.runs[i].text = para.runs[i].text[:-2] + new_p
                            para.runs[i+1].text = para.runs[i+1].text[1:]
                            replaced = True
                            break
                if not replaced and old_p in para.text:
                    para.text = para.text.replace(old_p, new_p)

    replace_in_paras(doc.paragraphs)
    def process_tables(tables):
        for table in tables:
            for row in table.rows:
                for cell in row.cells:
                    replace_in_paras(cell.paragraphs)
                    process_tables(cell.tables)
    process_tables(doc.tables)
    
    for section in doc.sections:
        for header in [section.header, section.first_page_header, section.even_page_header]:
            if header:
                replace_in_paras(header.paragraphs)
                process_tables(header.tables)
        for footer in [section.footer, section.first_page_footer, section.even_page_footer]:
            if footer:
                replace_in_paras(footer.paragraphs)
                process_tables(footer.tables)
    doc.save(filepath)

def replace_in_xlsx(filepath):
    wb = openpyxl.load_workbook(filepath)
    for sheet in wb.sheetnames:
        ws = wb[sheet]
        for row in ws.iter_rows():
            for cell in row:
                if cell.value and isinstance(cell.value, str):
                    if "IS-" in cell.value:
                        cell.value = cell.value.replace("IS-", "qpower-")
    wb.save(filepath)

def migrate_folder(folder):
    for root, dirs, files in os.walk(folder):
        for fname in files:
            if fname.endswith('.enc'):
                continue # We'll just replace the original and maybe they need to encrypt again
            fpath = os.path.join(root, fname)
            if fname.endswith('.docx'):
                replace_in_docx(fpath)
            elif fname.endswith('.xlsx'):
                replace_in_xlsx(fpath)
            
            if 'IS-' in fname:
                new_fname = fname.replace('IS-', 'qpower-', 1)
                os.rename(fpath, os.path.join(root, new_fname))

def migrate_json(filepath):
    with open(filepath, 'r', encoding='utf-8') as f:
        content = f.read()
    # Replace IS-
    content = content.replace("IS-", "qpower-")
    # Replace bare "IS" default
    content = content.replace('"doc_prefix": "IS"', '"doc_prefix": "qpower"')
    with open(filepath, 'w', encoding='utf-8') as f:
        f.write(content)

if __name__ == "__main__":
    base = "/Users/imitate/Desktop/iso_dev"
    migrate_folder(os.path.join(base, "files/marked"))
    migrate_folder(os.path.join(base, "examples"))
    migrate_json(os.path.join(base, "document_structure.json"))
    migrate_json(os.path.join(base, "questionnaire_config.json"))
    print("Migration complete.")
